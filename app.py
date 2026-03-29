# app.py — Spartan AI
import streamlit as st
import requests
from requests.auth import HTTPBasicAuth
import json
import time
import re
import io
import warnings

warnings.filterwarnings("ignore", message="Unverified HTTPS request")

from PIL import Image

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx as docx_module
except ImportError:
    docx_module = None

try:
    import pytesseract
    TESSERACT_OK = True
except ImportError:
    pytesseract = None
    TESSERACT_OK = False

# ── Config ────────────────────────────────────────────────────────────────────
NGROK_URL       = "https://ona-overcritical-extrinsically.ngrok-free.dev"
OLLAMA_CHAT_URL = f"{NGROK_URL}/api/chat"
USERNAME        = "dgeurts"
PASSWORD        = "thaidakar21"
OCR_CONFIG      = r"--oem 3 --psm 6"

MODEL_MAP = {
    "Assignment Generation": "spartan-generator",
    "Assignment Grader":     "spartan-grader",
    "AI Content Detector":   "spartan-detector",
    "Student Chatbot":       "spartan-student",
}

TOOL_META = {
    "Assignment Generation": {
        "tag": "GEN", "index": "01", "color": "#00c8ff", "icon": "✦",
        "desc": "Generate custom assignments, rubrics, and prompts for any subject or grade.",
    },
    "Assignment Grader": {
        "tag": "GRD", "index": "02", "color": "#38e8c0", "icon": "◈",
        "desc": "Upload student work and receive detailed rubric-aligned grading feedback.",
    },
    "AI Content Detector": {
        "tag": "DET", "index": "03", "color": "#a78bfa", "icon": "◉",
        "desc": "Analyze submissions for AI-generated content and plagiarism signals.",
    },
    "Student Chatbot": {
        "tag": "STU", "index": "04", "color": "#f472b6", "icon": "◎",
        "desc": "A responsible, curriculum-aware assistant to support student learning.",
    },
}

# ── Tag Parser ────────────────────────────────────────────────────────────────
OUTPUT_TEXT_RE = re.compile(r'\[output-text\](.*?)\[/output-text\]', re.DOTALL)
OUTPUT_FILE_RE = re.compile(
    r'\[output-file-(txt|md|pdf|docx)\](.*?)\[/output-file-(?:txt|md|pdf|docx)\]',
    re.DOTALL
)

def parse_ai_response(raw: str) -> list[dict]:
    segments = []
    cursor = 0
    all_matches = []
    for m in OUTPUT_TEXT_RE.finditer(raw): all_matches.append(("text", m))
    for m in OUTPUT_FILE_RE.finditer(raw):  all_matches.append(("file", m))
    all_matches.sort(key=lambda x: x[1].start())
    for kind, m in all_matches:
        gap = raw[cursor:m.start()].strip()
        if gap: segments.append({"type": "text", "content": gap})
        if kind == "text":
            c = m.group(1).strip()
            if c: segments.append({"type": "text", "content": c})
        else:
            ext, c = m.group(1).lower(), m.group(2).strip()
            if c: segments.append({"type": "file", "ext": ext, "content": c})
        cursor = m.end()
    tail = raw[cursor:].strip()
    if tail: segments.append({"type": "text", "content": tail})
    return segments

def make_download_bytes(content: str, ext: str) -> tuple[bytes, str]:
    if ext in ("txt", "md"): return content.encode("utf-8"), "text/plain"
    if ext == "pdf":          return content.encode("utf-8"), "text/plain"
    if ext == "docx":
        if docx_module:
            doc = docx_module.Document()
            for line in content.split("\n"): doc.add_paragraph(line)
            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            return buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return content.encode("utf-8"), "text/plain"
    return content.encode("utf-8"), "text/plain"

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="Spartan AI", layout="wide", initial_sidebar_state="expanded")

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=Space+Mono:wght@400;700&display=swap');

:root {
    --bg:      #030b18;
    --bg1:     #061226;
    --bg2:     #091833;
    --bg3:     #0c1e40;
    --blue:    #00b4ff;
    --blue-lo: rgba(0,180,255,0.08);
    --blue-bd: rgba(0,180,255,0.22);
    --blue-gl: rgba(0,180,255,0.30);
    --cyan:    #06e5d4;
    --cyan-lo: rgba(6,229,212,0.07);
    --cyan-bd: rgba(6,229,212,0.24);
    --txt:     rgba(200,228,255,0.92);
    --txt2:    rgba(100,148,210,0.65);
    --txt3:    rgba(55,95,160,0.52);
    --bdr:     rgba(0,140,220,0.14);
    --bdr2:    rgba(0,180,255,0.20);
}

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

html, body, [class*="css"], .stApp {
    background: var(--bg) !important;
    font-family: 'Space Grotesk', sans-serif !important;
    color: var(--txt) !important;
    -webkit-font-smoothing: antialiased !important;
}

.stApp {
    background:
        radial-gradient(ellipse 70% 45% at 55% -5%, rgba(0,110,255,0.07) 0%, transparent 60%),
        radial-gradient(ellipse 35% 30% at 95% 98%, rgba(6,229,212,0.05) 0%, transparent 55%),
        repeating-linear-gradient(0deg,   rgba(0,150,255,0.022) 0, rgba(0,150,255,0.022) 1px, transparent 1px, transparent 44px),
        repeating-linear-gradient(90deg,  rgba(0,150,255,0.022) 0, rgba(0,150,255,0.022) 1px, transparent 1px, transparent 44px),
        var(--bg) !important;
}

::-webkit-scrollbar { width: 3px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--blue-bd); border-radius: 3px; }

#MainMenu, footer, header,
div[data-testid="stDecoration"],
div[data-testid="stStatusWidget"],
.stDeployButton { display: none !important; }

/* ══════════════════ SIDEBAR ══════════════════ */
section[data-testid="stSidebar"] {
    background: rgba(3,11,24,0.95) !important;
    backdrop-filter: blur(20px) !important;
    -webkit-backdrop-filter: blur(20px) !important;
    border-right: 1px solid var(--bdr) !important;
    box-shadow: 2px 0 28px rgba(0,0,0,0.5) !important;
    width: 230px !important; min-width: 230px !important; max-width: 230px !important;
}
section[data-testid="stSidebar"] > div {
    padding: 0 !important;
    width: 230px !important; min-width: 230px !important; max-width: 230px !important;
    overflow-x: hidden !important;
}

.sb-brand {
    padding: 22px 16px 16px; border-bottom: 1px solid var(--bdr); position: relative;
}
.sb-brand::after {
    content: ''; position: absolute; bottom: 0; left: 0; right: 0; height: 1px;
    background: linear-gradient(90deg, rgba(0,180,255,0.35), transparent);
}
.sb-top { display: flex; align-items: center; gap: 11px; margin-bottom: 8px; }
.sb-icon {
    width: 34px; height: 34px; flex-shrink: 0;
    border: 1.5px solid var(--blue); border-radius: 8px;
    display: flex; align-items: center; justify-content: center;
    box-shadow: 0 0 12px rgba(0,180,255,0.22), inset 0 0 8px rgba(0,180,255,0.05);
}
.sb-icon span { font-family: 'Space Mono', monospace; font-size: 0.7rem; font-weight: 700; color: var(--blue); }
.sb-name { font-size: 0.9rem; font-weight: 600; color: var(--txt); letter-spacing: -0.01em; }
.sb-ver  { font-family: 'Space Mono', monospace; font-size: 0.44rem; color: var(--txt3); letter-spacing: 0.2em; text-transform: uppercase; margin-top: 1px; }
.sb-status {
    display: flex; align-items: center; gap: 7px; padding-left: 45px;
    font-family: 'Space Mono', monospace; font-size: 0.44rem; color: var(--txt3);
    letter-spacing: 0.14em; text-transform: uppercase;
}
.sb-dot {
    width: 5px; height: 5px; border-radius: 50%;
    background: var(--cyan); box-shadow: 0 0 5px var(--cyan);
    animation: pulse 2.5s ease-in-out infinite; flex-shrink: 0;
}
@keyframes pulse { 0%,100%{opacity:1;} 50%{opacity:0.3;} }

.sb-sec {
    padding: 14px 16px 4px;
    font-family: 'Space Mono', monospace; font-size: 0.44rem;
    letter-spacing: 0.24em; text-transform: uppercase; color: var(--txt3);
    display: flex; align-items: center; gap: 8px;
}
.sb-sec::after { content: ''; flex: 1; height: 1px; background: var(--bdr); }

div[data-testid="stSidebar"] .stButton {
    width: 100% !important; padding: 0 8px !important; margin: 0 !important;
}
div[data-testid="stSidebar"] .stButton > button {
    all: unset !important;
    display: flex !important; align-items: center !important;
    width: 100% !important;
    padding: 8px 10px !important; margin: 1px 0 !important;
    border-radius: 6px !important; border-left: 2px solid transparent !important;
    font-family: 'Space Grotesk', sans-serif !important;
    font-size: 0.78rem !important; font-weight: 400 !important; color: var(--txt2) !important;
    cursor: pointer !important;
    white-space: nowrap !important; overflow: hidden !important; text-overflow: ellipsis !important;
    transition: all 0.15s !important; box-sizing: border-box !important;
}
div[data-testid="stSidebar"] .stButton > button:hover {
    background: var(--blue-lo) !important; border-left-color: var(--blue) !important;
    color: var(--txt) !important; padding-left: 14px !important;
}

.sb-foot {
    padding: 12px 16px 16px; margin-top: 4px; border-top: 1px solid var(--bdr);
}
.sb-foot-line {
    font-family: 'Space Mono', monospace;
    font-size: 0.44rem; color: var(--txt3); letter-spacing: 0.1em; text-transform: uppercase; line-height: 2.2;
}

/* ══════════════════ MAIN ══════════════════ */
.main .block-container {
    max-width: 860px !important;
    padding: 0 2.5rem 0 !important;
}

/* ══════════════════ HOME ══════════════════ */
.mc-wrap { padding: 56px 0 40px; }
.mc-sys {
    font-family: 'Space Mono', monospace; font-size: 0.5rem;
    letter-spacing: 0.28em; text-transform: uppercase; color: var(--blue); opacity: 0.65;
    margin-bottom: 18px; display: flex; align-items: center; gap: 10px;
}
.mc-sys span { color: var(--blue); opacity: 1; margin-right: 2px; }
.mc-title {
    font-size: 3.1rem; font-weight: 700; line-height: 1;
    letter-spacing: -0.04em; color: var(--txt); margin-bottom: 16px;
}
.mc-title .dim { font-weight: 300; color: var(--txt2); }
.mc-title .ac  { color: var(--blue); }
.mc-subtitle {
    font-size: 0.88rem; color: var(--txt2); line-height: 1.72;
    max-width: 460px; font-weight: 300; margin-bottom: 48px;
}
.mc-rule { height: 1px; background: linear-gradient(90deg, var(--blue) 0%, transparent 55%); opacity: 0.22; margin-bottom: 36px; }

/* Module grid — all in ONE st.markdown call */
.mgrid {
    display: grid; grid-template-columns: 1fr 1fr; gap: 1px;
    background: var(--bdr); border: 1px solid var(--bdr); border-radius: 12px; overflow: hidden;
    margin-bottom: 48px;
    box-shadow: 0 8px 48px rgba(0,0,0,0.45), 0 0 0 1px rgba(0,120,255,0.05);
}
.mg {
    background: var(--bg1); padding: 28px 26px; position: relative; overflow: hidden;
    transition: background 0.18s;
}
.mg::before {
    content: ''; position: absolute; top: 0; left: 0; width: 3px; height: 100%;
    background: currentColor; opacity: 0; transition: opacity 0.18s;
}
.mg:hover { background: var(--bg2); }
.mg:hover::before { opacity: 0.55; }
.mg-num {
    font-family: 'Space Mono', monospace; font-size: 0.46rem;
    letter-spacing: 0.2em; text-transform: uppercase; color: var(--txt3); margin-bottom: 15px;
}
.mg-tag {
    display: inline-flex; align-items: center; gap: 5px;
    font-family: 'Space Mono', monospace; font-size: 0.52rem; font-weight: 700; letter-spacing: 0.1em;
    padding: 3px 9px; border-radius: 4px; background: rgba(0,0,0,0.25); border: 1px solid currentColor;
    margin-bottom: 13px; opacity: 0.82;
}
.mg-tag::before { content: ''; width: 4px; height: 4px; border-radius: 50%; background: currentColor; }
.mg-name { font-size: 0.95rem; font-weight: 600; color: var(--txt); margin-bottom: 8px; letter-spacing: -0.015em; }
.mg-desc { font-size: 0.73rem; color: var(--txt2); line-height: 1.65; font-weight: 300; }
.mg-sym  { position: absolute; right: 20px; bottom: 16px; font-size: 1.8rem; opacity: 0.05; color: currentColor; }

.mc-foot { display: flex; align-items: center; justify-content: space-between; border-top: 1px solid var(--bdr); padding-top: 18px; }
.mc-foot-l { font-family: 'Space Mono', monospace; font-size: 0.46rem; color: var(--txt3); letter-spacing: 0.1em; text-transform: uppercase; }
.mc-foot-r { font-family: 'Space Mono', monospace; font-size: 0.44rem; color: var(--blue); opacity: 0.6; border: 1px solid var(--bdr2); border-radius: 3px; padding: 2px 8px; letter-spacing: 0.14em; text-transform: uppercase; }

[data-testid="stHorizontalBlock"] [data-testid="stVerticalBlock"] .stButton,
[data-testid="stHorizontalBlock"] [data-testid="stVerticalBlock"] .stButton > button {
    display: none !important; height: 0 !important; padding: 0 !important; margin: 0 !important; overflow: hidden !important;
}

/* ══════════════════ TOOL HEADER ══════════════════ */
.th {
    padding: 26px 0 18px; border-bottom: 1px solid var(--bdr); margin-bottom: 24px; position: relative;
}
.th::after {
    content: ''; position: absolute; bottom: -1px; left: 0; width: 56px; height: 1px;
    background: var(--blue); box-shadow: 0 0 8px rgba(0,180,255,0.5);
}
.th-crumb {
    font-family: 'Space Mono', monospace; font-size: 0.46rem; letter-spacing: 0.18em; text-transform: uppercase;
    color: var(--txt3); margin-bottom: 11px; display: flex; align-items: center; gap: 7px;
}
.th-crumb .sep { color: var(--blue); opacity: 0.4; }
.th-crumb .cur { color: var(--blue); opacity: 0.72; }
.th-row { display: flex; align-items: flex-start; justify-content: space-between; }
.th-left { display: flex; align-items: center; gap: 14px; }
.th-badge {
    font-family: 'Space Mono', monospace; font-size: 0.54rem; font-weight: 700; letter-spacing: 0.1em;
    padding: 4px 10px; border-radius: 4px; border: 1px solid currentColor; flex-shrink: 0; opacity: 0.82;
}
.th-name { font-size: 1.08rem; font-weight: 600; color: var(--txt); letter-spacing: -0.02em; }
.th-desc { font-family: 'Space Mono', monospace; font-size: 0.53rem; color: var(--txt3); margin-top: 3px; line-height: 1.6; }
.th-live {
    display: flex; align-items: center; gap: 6px; flex-shrink: 0; padding-top: 3px;
    font-family: 'Space Mono', monospace; font-size: 0.44rem; color: var(--cyan);
    letter-spacing: 0.16em; text-transform: uppercase;
}
.th-live-dot { width: 5px; height: 5px; border-radius: 50%; background: var(--cyan); box-shadow: 0 0 5px var(--cyan); animation: pulse 2.5s ease-in-out infinite; }

/* ══════════════════ CHAT MESSAGES ══════════════════ */
.stChatMessage {
    background: transparent !important; border: none !important; padding: 5px 0 !important; gap: 10px !important;
}
div[data-testid="chatAvatarIcon-user"],
div[data-testid="chatAvatarIcon-assistant"] {
    width: 28px !important; height: 28px !important;
    border-radius: 7px !important; overflow: hidden !important; flex-shrink: 0 !important;
}
div[data-testid="chatAvatarIcon-user"] img, div[data-testid="chatAvatarIcon-assistant"] img,
div[data-testid="chatAvatarIcon-user"] svg, div[data-testid="chatAvatarIcon-assistant"] svg {
    filter: saturate(0) brightness(0.35) !important; width: 100% !important; height: 100% !important;
}
div[data-testid="chatAvatarIcon-user"]      { background: var(--blue-lo) !important; border: 1px solid var(--blue-bd) !important; }
div[data-testid="chatAvatarIcon-assistant"] { background: var(--cyan-lo) !important; border: 1px solid var(--cyan-bd) !important; }

div[data-testid="stChatMessageUser"] { flex-direction: row-reverse !important; }
div[data-testid="stChatMessageUser"] > div[data-testid="stChatMessageContent"] {
    background: var(--blue-lo) !important; backdrop-filter: blur(10px) !important;
    border: 1px solid var(--blue-bd) !important;
    border-radius: 14px 3px 14px 14px !important;
    padding: 12px 16px !important; max-width: 70% !important;
    font-size: 0.88rem !important; line-height: 1.7 !important; color: var(--txt) !important;
    box-shadow: 0 2px 12px rgba(0,0,0,0.25) !important;
}
div[data-testid="stChatMessageAssistant"] > div[data-testid="stChatMessageContent"] {
    background: rgba(6,18,38,0.78) !important; backdrop-filter: blur(12px) !important;
    border: 1px solid var(--bdr2) !important; border-left: 2px solid var(--cyan) !important;
    border-radius: 3px 14px 14px 14px !important;
    padding: 13px 18px !important; max-width: 88% !important;
    font-size: 0.88rem !important; line-height: 1.76 !important; color: rgba(210,232,255,0.93) !important;
    box-shadow: 0 2px 16px rgba(0,0,0,0.22) !important;
}

/* ══════════════════ FILE CHIP ══════════════════ */
.fc {
    display: inline-flex; align-items: center; gap: 8px;
    background: rgba(0,180,255,0.07); border: 1px solid var(--blue-bd);
    border-radius: 8px; padding: 8px 14px;
    margin-bottom: 8px;
    font-family: 'Space Mono', monospace;
    box-shadow: 0 0 12px rgba(0,180,255,0.07);
}
.fc-icon { font-size: 0.9rem; flex-shrink: 0; }
.fc-body { display: flex; flex-direction: column; gap: 1px; }
.fc-name { font-size: 0.62rem; color: var(--txt); font-weight: 700; letter-spacing: 0.02em; }
.fc-sub  { font-size: 0.5rem;  color: var(--cyan); letter-spacing: 0.1em; text-transform: uppercase; opacity: 0.85; }

/* ══════════════════ UPLOAD PANEL ══════════════════ */
.up-panel {
    background: var(--bg2);
    border: 1px solid var(--bdr2);
    border-radius: 12px;
    padding: 20px 22px 16px;
    margin-bottom: 10px;
    box-shadow: 0 4px 28px rgba(0,0,0,0.35), 0 0 0 1px rgba(0,150,255,0.04);
    position: relative;
}
.up-title {
    font-family: 'Space Mono', monospace; font-size: 0.54rem;
    letter-spacing: 0.18em; text-transform: uppercase; color: var(--blue); opacity: 0.75;
    margin-bottom: 12px; display: flex; align-items: center; gap: 8px;
}
.up-title::before { content: '//'; font-weight: 700; color: var(--blue); }
div[data-testid="stFileUploader"] > div {
    background: rgba(6,18,38,0.7) !important;
    border: 1px dashed var(--bdr2) !important;
    border-radius: 8px !important;
    padding: 16px 14px !important;
    transition: border-color 0.18s, background 0.18s !important;
}
div[data-testid="stFileUploader"] > div:hover {
    background: var(--blue-lo) !important;
    border-color: var(--blue-bd) !important;
    box-shadow: 0 0 14px rgba(0,180,255,0.07) !important;
}
div[data-testid="stFileUploader"] label,
div[data-testid="stFileUploader"] small {
    font-family: 'Space Mono', monospace !important;
    color: var(--txt2) !important; font-size: 0.6rem !important;
}
div[data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] {
    font-family: 'Space Mono', monospace !important;
    font-size: 0.56rem !important; color: var(--cyan) !important;
}

/* ══════════════════ FILE DOWNLOAD CARD ══════════════════ */
.fcard {
    background: var(--bg2); border: 1px solid var(--bdr2); border-left: 2px solid var(--blue);
    border-radius: 10px; padding: 18px 20px; margin: 8px 0;
    box-shadow: 0 4px 24px rgba(0,0,0,0.28);
}
.fcard-hd { display: flex; align-items: center; gap: 12px; margin-bottom: 12px; }
.fcard-icon {
    width: 36px; height: 36px; border-radius: 6px; flex-shrink: 0;
    background: var(--blue-lo); border: 1px solid var(--blue-bd);
    display: flex; align-items: center; justify-content: center;
    font-family: 'Space Mono', monospace; font-size: 0.58rem; font-weight: 700; color: var(--blue);
}
.fcard-name { font-size: 0.84rem; font-weight: 600; color: var(--txt); }
.fcard-meta { font-family: 'Space Mono', monospace; font-size: 0.5rem; color: var(--txt3); letter-spacing: 0.05em; margin-top: 2px; }
.fcard-pre  {
    background: rgba(2,8,20,0.72); border: 1px solid var(--bdr); border-radius: 6px;
    padding: 10px 14px; margin-bottom: 13px;
    font-family: 'Space Mono', monospace; font-size: 0.66rem; color: var(--txt2);
    line-height: 1.6; white-space: pre-wrap; word-break: break-word;
    max-height: 100px; overflow: hidden; position: relative;
}
.fcard-pre::after {
    content: ''; position: absolute; bottom: 0; left: 0; right: 0; height: 28px;
    background: linear-gradient(transparent, rgba(2,8,20,0.9)); border-radius: 0 0 6px 6px;
}
.fcard .stDownloadButton > button {
    all: unset !important;
    display: inline-flex !important; align-items: center !important; gap: 7px !important;
    font-family: 'Space Mono', monospace !important; font-size: 0.56rem !important; letter-spacing: 0.1em !important;
    color: #030b18 !important; background: var(--blue) !important;
    border-radius: 6px !important; padding: 8px 16px !important; cursor: pointer !important; font-weight: 700 !important;
    box-shadow: 0 0 18px rgba(0,180,255,0.28) !important; transition: box-shadow 0.15s, opacity 0.15s !important;
}
.fcard .stDownloadButton > button:hover { box-shadow: 0 0 28px rgba(0,180,255,0.48) !important; opacity: 0.9 !important; }

/* ══════════════════ FILE GEN CARD ══════════════════ */
.fgen {
    background: var(--bg2); border: 1px solid var(--cyan-bd); border-left: 2px solid var(--cyan);
    border-radius: 10px; padding: 15px 18px; display: flex; align-items: center; gap: 14px; margin: 8px 0;
}
.fgen-icon { width: 36px; height: 36px; border-radius: 6px; flex-shrink: 0; background: var(--cyan-lo); border: 1px solid var(--cyan-bd); display: flex; align-items: center; justify-content: center; }
.fgen-icon svg { width: 15px; height: 15px; }
.fgen-title { font-size: 0.82rem; font-weight: 600; color: var(--txt); margin-bottom: 4px; }
.fgen-sub { font-family: 'Space Mono', monospace; font-size: 0.5rem; color: var(--cyan); letter-spacing: 0.08em; display: flex; align-items: center; gap: 7px; }
.fdots { display: flex; gap: 3px; }
.fdots span { display: block; width: 3px; height: 3px; border-radius: 50%; background: var(--cyan); animation: blink 1.1s ease-in-out infinite both; }
.fdots span:nth-child(2){animation-delay:0.2s;} .fdots span:nth-child(3){animation-delay:0.4s;}

/* ══════════════════ THINKING ══════════════════ */
.thinking { display: inline-flex; align-items: center; gap: 10px; padding: 6px 0; }
.tdots { display: flex; gap: 5px; }
.tdots span { display: block; width: 6px; height: 6px; border-radius: 50%; background: var(--blue); box-shadow: 0 0 6px rgba(0,180,255,0.4); animation: blink 1.1s ease-in-out infinite both; }
.tdots span:nth-child(2){animation-delay:0.2s;} .tdots span:nth-child(3){animation-delay:0.4s;}
@keyframes blink { 0%,80%,100%{opacity:0.1;transform:scale(0.6);} 40%{opacity:1;transform:scale(1);} }
.thinking-lbl { font-family: 'Space Mono', monospace; font-size: 0.52rem; color: var(--txt3); letter-spacing: 0.2em; text-transform: uppercase; }

/* ══════════════════ STICKY BOTTOM BAR ══════════════════ */
/* Target the horizontal block that contains the chat input */
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) {
    position: sticky !important;
    bottom: 0 !important;
    z-index: 100 !important;
    gap: 8px !important;
    padding: 14px 0 28px !important;
    background: linear-gradient(to top, var(--bg) 60%, transparent) !important;
}

/* Column sizing for the bottom bar */
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) > div:nth-child(1),
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) > div:nth-child(2) {
    flex: 0 0 44px !important; min-width: 44px !important; max-width: 44px !important; padding: 0 !important;
}

/* Bottom bar icon buttons */
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) > div:nth-child(1) .stButton > button,
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) > div:nth-child(2) .stButton > button {
    all: unset !important;
    width: 44px !important; height: 44px !important;
    display: flex !important; align-items: center !important; justify-content: center !important;
    background: rgba(6,18,38,0.92) !important; backdrop-filter: blur(14px) !important;
    border: 1px solid var(--bdr2) !important; border-radius: 10px !important;
    font-size: 1.05rem !important; color: var(--txt3) !important;
    cursor: pointer !important; transition: all 0.15s !important; box-sizing: border-box !important;
}
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) > div:nth-child(1) .stButton > button:hover,
div[data-testid="stHorizontalBlock"]:has(div[data-testid="stChatInput"]) > div:nth-child(2) .stButton > button:hover {
    background: var(--blue-lo) !important; border-color: var(--blue-bd) !important;
    color: var(--blue) !important; box-shadow: 0 0 14px rgba(0,180,255,0.12) !important;
}

/* ══════════════════ CHAT INPUT ══════════════════ */
div[data-testid="stChatInput"] {
    background: rgba(6,18,38,0.94) !important;
    backdrop-filter: blur(18px) !important; -webkit-backdrop-filter: blur(18px) !important;
    border: 1px solid var(--bdr2) !important; border-radius: 12px !important;
    box-shadow: 0 4px 28px rgba(0,0,0,0.35), inset 0 1px 0 rgba(0,180,255,0.04) !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
div[data-testid="stChatInput"]:focus-within {
    border-color: rgba(0,180,255,0.4) !important;
    box-shadow: 0 4px 28px rgba(0,0,0,0.35), 0 0 0 3px rgba(0,180,255,0.06), 0 0 20px rgba(0,180,255,0.07) !important;
}
div[data-testid="stChatInput"] textarea {
    background: transparent !important; color: var(--txt) !important;
    font-family: 'Space Grotesk', sans-serif !important; font-size: 0.9rem !important;
    caret-color: var(--blue) !important;
}
div[data-testid="stChatInput"] textarea::placeholder {
    color: var(--txt3) !important; font-family: 'Space Mono', monospace !important;
    font-size: 0.75rem !important; letter-spacing: 0.04em !important;
}
div[data-testid="stChatInput"] button {
    background: var(--blue) !important; border: none !important; border-radius: 7px !important;
    width: 30px !important; height: 30px !important;
    display: flex !important; align-items: center !important; justify-content: center !important;
    box-shadow: 0 0 14px rgba(0,180,255,0.3) !important; transition: box-shadow 0.15s !important;
    flex-shrink: 0 !important; margin: auto 6px auto 0 !important; padding: 0 !important;
}
div[data-testid="stChatInput"] button:hover { box-shadow: 0 0 22px rgba(0,180,255,0.5) !important; }
div[data-testid="stChatInput"] button svg { width: 13px !important; height: 13px !important; }
div[data-testid="stChatInput"] button svg path,
div[data-testid="stChatInput"] button svg rect { fill: #030b18 !important; stroke: none !important; }

/* ══════════════════ ALERTS ══════════════════ */
div[data-testid="stAlert"] {
    background: var(--blue-lo) !important; border: 1px solid var(--blue-bd) !important;
    border-radius: 8px !important; font-family: 'Space Mono', monospace !important; font-size: 0.68rem !important;
}
</style>
""", unsafe_allow_html=True)

# ── Session state ─────────────────────────────────────────────────────────────
for k, v in [("mode","Home"),("messages",[]),("pending_ocr_text",None),
              ("uploaded_file_name",None),("show_upload",False)]:
    if k not in st.session_state:
        st.session_state[k] = v

def go_to_tool(name: str):
    st.session_state.mode               = name
    st.session_state.messages           = [{"role":"assistant","content":"[output-text]System online. How can I assist you today?[/output-text]"}]
    st.session_state.pending_ocr_text   = None
    st.session_state.uploaded_file_name = None
    st.session_state.show_upload        = False

def render_assistant_message(raw: str, msg_index: int):
    segments = parse_ai_response(raw)
    if not segments: st.markdown(raw); return
    for si, seg in enumerate(segments):
        if seg["type"] == "text":
            st.markdown(seg["content"])
        else:
            ext, content = seg["ext"], seg["content"]
            fname   = f"Spartan-Assignment.{ext}"
            preview = content[:380]
            size_kb = round(len(content.encode()) / 1024, 1)
            lbl = {"txt":"TXT","md":"MD","pdf":"PDF","docx":"DOCX"}.get(ext, ext.upper())
            st.markdown(f"""
<div class="fcard">
  <div class="fcard-hd">
    <div class="fcard-icon">{lbl}</div>
    <div><div class="fcard-name">{fname}</div><div class="fcard-meta">{size_kb} KB &middot; OUTPUT READY</div></div>
  </div>
  <div class="fcard-pre">{preview}</div>
</div>""", unsafe_allow_html=True)
            fb, mime = make_download_bytes(content, ext)
            st.download_button(f"↓  DOWNLOAD  {fname}", data=fb, file_name=fname, mime=mime,
                               key=f"dl_{msg_index}_{si}_{ext}")

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
<div class="sb-brand">
  <div class="sb-top">
    <div class="sb-icon"><span>S</span></div>
    <div><div class="sb-name">Spartan AI</div><div class="sb-ver">System v1.0</div></div>
  </div>
  <div class="sb-status"><div class="sb-dot"></div>All systems operational</div>
</div>
<div class="sb-sec">Navigate</div>
""", unsafe_allow_html=True)

    if st.button("⌂  Home", key="sb_home"):
        st.session_state.mode = "Home"
        st.session_state.messages = []
        st.session_state.pending_ocr_text = None
        st.session_state.uploaded_file_name = None
        st.session_state.show_upload = False
        st.rerun()

    st.markdown('<div class="sb-sec">Modules</div>', unsafe_allow_html=True)

    for tool_name, tm in TOOL_META.items():
        if st.button(f"{tm['icon']}  {tm['index']}  {tool_name}", key=f"sb_{tool_name}"):
            go_to_tool(tool_name)
            st.rerun()

    st.markdown("""
<div class="sb-foot">
  <div class="sb-foot-line">Senior Project &middot; 2025</div>
  <div class="sb-foot-line">Dallin Geurts</div>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# HOME
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.mode == "Home":
    mods = ""
    for name, tm in TOOL_META.items():
        mods += f"""
<div class="mg" style="color:{tm['color']};">
  <div class="mg-num">Module {tm['index']}</div>
  <div class="mg-tag">{tm['tag']}</div>
  <div class="mg-name">{name}</div>
  <div class="mg-desc">{tm['desc']}</div>
  <div class="mg-sym">{tm['icon']}</div>
</div>"""

    st.markdown(f"""
<div class="mc-wrap">
  <div class="mc-sys"><span>//</span> Educational Intelligence Platform</div>
  <div class="mc-title"><span class="dim">SPARTAN</span> <span class="ac">AI</span></div>
  <div class="mc-subtitle">Four specialized AI modules for educators and students —<br>built for speed, transparency, and trust.</div>
</div>
<div class="mc-rule"></div>
<div class="mgrid">{mods}</div>
<div class="mc-foot">
  <div class="mc-foot-l">Spartan AI &middot; Dallin Geurts &middot; 2025</div>
  <div class="mc-foot-r">v1.0 &middot; Ready</div>
</div>
<div style="height:32px;"></div>
""", unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    for i, (name, _) in enumerate(TOOL_META.items()):
        with (c1 if i % 2 == 0 else c2):
            if st.button(name, key=f"home_{name}"):
                go_to_tool(name)
                st.rerun()
    st.stop()

# ══════════════════════════════════════════════════════════════════════════════
# TOOL PAGE
# ══════════════════════════════════════════════════════════════════════════════
tool  = st.session_state.mode
model = MODEL_MAP[tool]
tm    = TOOL_META[tool]

st.markdown(f"""
<div class="th">
  <div class="th-crumb">Spartan AI <span class="sep">/</span> <span class="cur">{tm['tag']}</span></div>
  <div class="th-row">
    <div class="th-left">
      <div class="th-badge" style="color:{tm['color']};border-color:{tm['color']}44;">{tm['tag']}</div>
      <div><div class="th-name">{tool}</div><div class="th-desc">{tm['desc']}</div></div>
    </div>
    <div class="th-live"><div class="th-live-dot"></div>Online</div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Chat history ──────────────────────────────────────────────────────────────
for i, msg in enumerate(st.session_state.messages):
    avatar = "🤖" if msg["role"] == "assistant" else "👤"
    with st.chat_message(msg["role"], avatar=avatar):
        if msg["role"] == "assistant":
            render_assistant_message(msg["content"], i)
        else:
            st.markdown(msg.get("display_text", msg["content"]))

# Spacer so content clears the sticky bar
st.markdown("<div style='height:10px;'></div>", unsafe_allow_html=True)

# ── File attachment chip (visible when a file is staged) ──────────────────────
if st.session_state.pending_ocr_text:
    fname_display = st.session_state.uploaded_file_name or "file"
    st.markdown(f"""
<div class="fc">
  <div class="fc-icon">📎</div>
  <div class="fc-body">
    <div class="fc-name">{fname_display}</div>
    <div class="fc-sub">Attached · will send with next message</div>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Upload panel (shown when attach button toggled) ───────────────────────────
if st.session_state.show_upload:
    st.markdown("""
<div class="up-panel">
  <div class="up-title">Attach File</div>
</div>
""", unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "Drop a file or click to browse — PDF, DOCX, TXT, or image",
        type=["pdf","docx","txt","png","jpg","jpeg","gif","bmp","tiff"],
        label_visibility="visible",
        key="file_uploader_panel",
    )
    if uploaded_file is not None and uploaded_file.name != st.session_state.uploaded_file_name:
        with st.spinner("Reading file…"):
            extracted = ""
            ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
            try:
                if ext == "pdf":
                    if PyPDF2:
                        reader = PyPDF2.PdfReader(uploaded_file)
                        for page in reader.pages:
                            t = page.extract_text()
                            if t: extracted += t + "\n"
                    else: extracted = "(PyPDF2 not installed)"
                elif ext == "docx":
                    if docx_module:
                        d = docx_module.Document(uploaded_file)
                        for p in d.paragraphs: extracted += p.text + "\n"
                    else: extracted = "(python-docx not installed)"
                elif ext == "txt":
                    extracted = uploaded_file.read().decode("utf-8", errors="ignore")
                elif ext in ["png","jpg","jpeg","gif","bmp","tiff"]:
                    if TESSERACT_OK:
                        img = Image.open(uploaded_file).convert("RGB")
                        extracted = pytesseract.image_to_string(img, config=OCR_CONFIG)
                    else: extracted = "(pytesseract not installed)"
                else: extracted = "(Unsupported file type)"
                extracted = extracted.strip() or "(No readable text found)"
                st.session_state.pending_ocr_text   = extracted
                st.session_state.uploaded_file_name = uploaded_file.name
                st.session_state.show_upload         = False
                st.rerun()
            except Exception as e:
                st.error(f"File read error: {e}")
else:
    uploaded_file = None

# ── BOTTOM BAR — sticky: [new] [attach] [input] ──────────────────────────────
bc_new, bc_attach, bc_input = st.columns([1, 1, 14], gap="small")

with bc_new:
    if st.button("↺", key="new_chat", help="New session"):
        go_to_tool(tool)
        st.rerun()

with bc_attach:
    icon = "✕" if st.session_state.show_upload else "📎"
    if st.button(icon, key="attach_btn", help="Attach file"):
        st.session_state.show_upload = not st.session_state.show_upload
        st.rerun()

with bc_input:
    user_input = st.chat_input("> _  message Spartan AI…")

# ── Handle user input ─────────────────────────────────────────────────────────
if user_input:
    ocr = st.session_state.pending_ocr_text
    if ocr:
        api_content = f"[input-image-text]{ocr}[/input-image-text]\n[output-text]{user_input}[/output-text]"
        st.session_state.pending_ocr_text   = None
        st.session_state.uploaded_file_name = None
    else:
        api_content = f"[output-text]{user_input}[/output-text]"

    st.session_state.messages.append({"role":"user","content":api_content,"display_text":user_input})
    with st.chat_message("user", avatar="👤"):
        st.markdown(user_input)

    full_response = ""

    with st.chat_message("assistant", avatar="🤖"):
        thinking_slot = st.empty()
        resp_slot     = st.empty()
        thinking_slot.markdown("""
<div class="thinking">
  <div class="tdots"><span></span><span></span><span></span></div>
  <div class="thinking-lbl">Processing</div>
</div>""", unsafe_allow_html=True)

        try:
            payload = {
                "model": model,
                "messages": [{"role":m["role"],"content":m["content"]} for m in st.session_state.messages],
                "stream": True,
            }
            ANY_OPEN_RE  = re.compile(r'\[(output-text|output-file-(?:txt|md|pdf|docx))\]')
            FILE_OPEN_RE = re.compile(r'\[output-file-(?:txt|md|pdf|docx)\]')
            FILE_GEN_HTML = """
<div class="fgen">
  <div class="fgen-icon">
    <svg viewBox="0 0 16 16" fill="none"><path d="M3 2h7l4 4v9H3V2z" stroke="#06e5d4" stroke-width="1.2" stroke-linejoin="round"/><path d="M10 2v4h4" stroke="#06e5d4" stroke-width="1.2" stroke-linejoin="round"/><path d="M5 9h6M5 11.5h4" stroke="#06e5d4" stroke-width="1" stroke-linecap="round"/></svg>
  </div>
  <div>
    <div class="fgen-title">Generating output file…</div>
    <div class="fgen-sub">Writing document <div class="fdots"><span></span><span></span><span></span></div></div>
  </div>
</div>"""

            stream_state = "waiting"
            gen_slot = None
            active_slot = resp_slot

            with requests.post(
                OLLAMA_CHAT_URL, json=payload,
                auth=HTTPBasicAuth(USERNAME, PASSWORD),
                timeout=600, verify=False, stream=True,
            ) as r:
                r.raise_for_status()
                for line in r.iter_lines():
                    if not line: continue
                    token = json.loads(line).get("message", {}).get("content", "")
                    full_response += token

                    if stream_state == "waiting":
                        thinking_slot.empty()
                        m = ANY_OPEN_RE.search(full_response)
                        if m:
                            if m.group(1) == "output-text":
                                stream_state = "text"
                                inner = full_response[m.end():]
                                live  = re.sub(r'\[/output-text\].*', '', inner).strip()
                                if live: active_slot.markdown(live + "▌", unsafe_allow_html=True)
                            else:
                                stream_state = "file"
                                pre = OUTPUT_TEXT_RE.sub(r'\1', full_response[:m.start()]).strip()
                                if pre: active_slot.markdown(pre)
                                gen_slot = st.empty()
                                gen_slot.markdown(FILE_GEN_HTML, unsafe_allow_html=True)
                    elif stream_state == "text":
                        m2 = re.search(r'\[output-text\](.*?)(?=\[/output-text\]|\[output-file-)', full_response, re.DOTALL)
                        if m2:
                            live = m2.group(1).strip()
                        else:
                            start = full_response.rfind('[output-text]')
                            live  = full_response[start + len('[output-text]'):].strip()
                            live  = re.sub(r'\[/?$|\[/output', '', live)
                        active_slot.markdown(live + "▌", unsafe_allow_html=True)
                        if FILE_OPEN_RE.search(full_response):
                            stream_state = "file"
                            active_slot.markdown(live)
                            gen_slot = st.empty()
                            gen_slot.markdown(FILE_GEN_HTML, unsafe_allow_html=True)
                    time.sleep(0.01)

                thinking_slot.empty(); resp_slot.empty()
                if gen_slot: gen_slot.empty()
                render_assistant_message(full_response, len(st.session_state.messages))

        except Exception:
            thinking_slot.empty()
            full_response = ""
            resp_slot.markdown(
                "<span style='font-family:Space Mono,monospace;font-size:0.65rem;"
                "color:rgba(255,65,65,0.7);letter-spacing:0.07em;'>"
                "ERR // Connection failed — check server status.</span>",
                unsafe_allow_html=True,
            )

    if full_response:
        st.session_state.messages.append({"role":"assistant","content":full_response,"display_text":full_response})
